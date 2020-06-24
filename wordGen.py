# 生成word, 用于代入每个人的个人信息
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
import os
source = r'xxxx'  # 生成的word存储位置


def genWord(name, paper):
    doc = Document()
    doc.styles['Normal'].font.name = u'宋体'   # Normal表示是设置默认的正文格式
    doc.styles["Normal"].font.size = Pt(12)  # 正文字号12（小四）
    # doc.paragraphs.paragraph_format.line_spacing = 1.25
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 第一标题
    p = doc.add_paragraph()
    run = p.add_run('题目1')
    run.font.size = Pt(16)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 第二标题
    p = doc.add_paragraph()
    run = p.add_run('题目2')
    run.font.size = Pt(20)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 抬头
    p = doc.add_paragraph()
    p.add_run('尊敬的'+name+':')
    # 问候
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    p.add_run('    您好！')
    # 正文1段
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    para1 = '    您撰写的《'
    para2 = '》,经组委会审定符合会议征稿要求，已被录用，热忱欢迎您届时出席会议。'
    p.add_run(para1 + paper + para2)
    # 正文2段
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    para = '    本次论坛定于2020年x月x日上午8:30～12:00，下午1:30～5:30举行，举办方式为线上。请您扫描文末二维码入群方便接收参会相关消息。如有任何疑问也欢迎您电话联系我们：'
    p.add_run(para)
    # 联系方式
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    p.add_run('xx：xxxxx')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    p.add_run('xx：xxxxx')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    p.add_run('xx：xxxxx')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    p.add_run('xx：xxxx')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 二维码
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture('weChat.png', width=Cm(5.61))
    run = p.add_run(19*' ')
    run.add_picture('weChat2.png', width=Cm(5.61))
    path = os.path.join(source, name + '.docx')
    # 有的人提交2个论文,通过这个检测
    if os.path.exists(path):
        print(name+'wrong')
    # 保存,
    doc.save(path)



# genWord('XX', 'XXX')
